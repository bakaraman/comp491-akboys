import os
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Generate UPDATED Gantt chart PNG ─────────────────────────────
_gantt_path = "/tmp/gantt_chart_updated.png"

def _make_gantt():
    fig, ax = plt.subplots(figsize=(12, 5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("#f8f8f8")

    tasks = [
        ("WP1: Research & Setup",              1,  2, "#4e79a7"),
        ("WP2: Core Engine & Web Platform",    3,  6, "#f28e2b"),
        ("WP3: Multiplayer & Persistence",     7, 10, "#59a14f"),
        ("WP4: Advanced Features & Deploy",    9, 12, "#e15759"),
        ("WP5: Testing & Polish",             13, 15, "#76b7b2"),
    ]
    milestones = [
        ("M1.1 Demo Ready",          2,  0),
        ("M2.2 Web App Live",        6,  1),
        ("M3.2 Multiplayer Co-op",  10,  2),
        ("M4.2 Deployed",           12,  3),
        ("M5.1 Poster",             14,  4),
        ("M5.2 Final Demo",         15,  4),
    ]

    # Mark completed WPs
    completed_weeks = 6  # WP1 + WP2 done

    y_pos = list(range(len(tasks) - 1, -1, -1))

    for i, (label, start, end, color) in enumerate(tasks):
        y = y_pos[i]
        ax.barh(y, end - start + 1, left=start - 0.5, height=0.55,
                color=color, alpha=0.85, edgecolor="white", linewidth=1.2)
        ax.text(start - 0.5 + (end - start + 1) / 2, y,
                f"W{start}–W{end}", ha="center", va="center",
                fontsize=8, color="white", fontweight="bold")

    for label, week, task_i in milestones:
        y = y_pos[task_i]
        marker = "D" if week > completed_weeks else "s"
        color = "#333333" if week > completed_weeks else "#2d8a2d"
        ax.plot(week, y, marker=marker, color=color, markersize=7, zorder=5)
        ax.text(week, y + 0.33, label, ha="center", va="bottom",
                fontsize=6.5, color=color, fontstyle="italic")

    # Progress line
    ax.axvline(x=completed_weeks + 0.5, color="#2d8a2d", linestyle="-", linewidth=2, alpha=0.7)
    ax.text(completed_weeks + 0.5, len(tasks) - 0.3, "← We are here (W7)", ha="left",
            fontsize=8, color="#2d8a2d", fontweight="bold")

    ax.set_xlim(0.5, 15.5)
    ax.set_xticks(range(1, 16))
    ax.set_xticklabels([f"W{i}" for i in range(1, 16)], fontsize=8)
    ax.set_yticks(y_pos)
    ax.set_yticklabels([t[0] for t in tasks], fontsize=9)
    ax.xaxis.grid(True, linestyle="--", alpha=0.5, color="#cccccc")
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.axvline(x=15, color="#e15759", linestyle=":", linewidth=1.5, alpha=0.6)
    ax.axvline(x=13, color="#76b7b2", linestyle=":", linewidth=1.5, alpha=0.6)
    ax.set_xlabel("Week", fontsize=9, labelpad=6)
    ax.set_title("Project Schedule (Updated) — COMP 491 Spring 2026",
                 fontsize=11, fontweight="bold", pad=10)
    plt.tight_layout(pad=1.2)
    plt.savefig(_gantt_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()

_make_gantt()

doc = Document("../lecture-slides/COMP491_Proposal_Template.docx")

# ── helpers ─────────────────────────────────────────────────────
def set_para(para, text, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    for run in para.runs:
        run.text = ""
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run()
    r.text = text
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = align

def set_cell(cell, text, size=10, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    r.text = text
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Times New Roman"

def add_row_to_table(table, values, size=10):
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell(row.cells[i], val, size=size)
    return row

def insert_para_after(para, text, style_name=None, size=12, bold=False,
                       italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before=0, space_after=6):
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    idx = list(doc.paragraphs).index(para) + 1
    new_para = doc.paragraphs[idx]
    if style_name:
        try:
            new_para.style = doc.styles[style_name]
        except Exception:
            pass
    new_para.paragraph_format.space_before = Pt(space_before)
    new_para.paragraph_format.space_after = Pt(space_after)
    new_para.alignment = align
    r = new_para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    return new_para

def find_para(fragment):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None

# ── 1. TITLE PAGE ────────────────────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip() == "PROJECT TITLE":
        set_para(p, "Text-Based Adventure with LLMs",
                 size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    elif p.text.strip() == "Fall 2025":
        set_para(p, "Spring 2026 (Updated)", align=WD_ALIGN_PARAGRAPH.CENTER)
    elif p.text.strip() == "Project Advisor:":
        set_para(p, "Project Advisor: Baris Akgun")

# ── 2. PARTICIPANT TABLE ─────────────────────────────────────────
members = [
    ("Kadir Yigit Ozcelik", "79975", "kadirozcelik21@ku.edu.tr", "+90 551 215 64 74"),
    ("Serdar Yengil",       "80232", "syengil21@ku.edu.tr",      "+90 545 775 08 63"),
    ("Batuhan Karaman",     "79791", "bkaraman21@ku.edu.tr",     "+90 534 878 90 16"),
    ("Ata Berke Goktekin",  "80277", "agoktekin21@ku.edu.tr",    "+90 538 064 49 44"),
]
for table in doc.tables:
    if table.rows[0].cells[0].text.strip() == "Name":
        for row_i, m in enumerate(members):
            if row_i + 1 < len(table.rows):
                for col_i, val in enumerate(m):
                    set_cell(table.rows[row_i + 1].cells[col_i], val)
        break

# ── 3. ABSTRACT ──────────────────────────────────────────────────
for p in doc.paragraphs:
    if "A brief summary of your project." in p.text:
        set_para(p, "")
    elif "Summarize the motivation" in p.text:
        set_para(p,
            "This project is a web-based platform for AI-powered text adventure games. "
            "Players explore story worlds using natural language, with a Large Language Model (LLM) acting as "
            "the narrator. The system uses a TypeScript monorepo architecture with a Next.js 15 frontend, "
            "Express.js backend, and OpenAI GPT-5.4 for streaming narration. Six handcrafted game scenarios "
            "(Noir, Haunted Manor, Space Station, Pirate, Western, Cyberpunk) are available, each with 5 rooms, "
            "3 NPCs, and 5 collectible items forming evidence chains. The platform features real-time streaming "
            "narration via Server-Sent Events (SSE), AI-generated action suggestions via GPT-5-nano, and "
            "session persistence. Planned features include Socket.IO-based multiplayer co-op, Firebase Auth "
            "for user identity, AI-generated scene visuals via gpt-image-1.5, and deployment on "
            "Google Cloud Run + Firebase Hosting."
        )

# ── 4. TABLE OF CONTENTS ─────────────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip() == "Table of contents":
        set_para(p, "Table of Contents", bold=True, size=12,
                 align=WD_ALIGN_PARAGRAPH.LEFT)

# ── 5. INTRODUCTION ──────────────────────────────────────────────
intro_text = (
    "Text-based adventure games, also known as interactive fiction, are one of the oldest forms of digital "
    "entertainment. Pioneered by Colossal Cave Adventure (1976) and Zork (1977), these games let players explore "
    "narratives through natural language. Despite their long history, the genre has been limited by hand-authored "
    "content: every room, NPC, and puzzle had to be written by a developer in advance.\n\n"
    "The emergence of Large Language Models (LLMs) opens a new chapter for this genre. AI Dungeon (2019) "
    "demonstrated that LLMs can generate coherent narratives on demand, but purely LLM-driven games suffer from "
    "a fundamental problem: the model has no concept of rules or game state. It can hallucinate items, contradict "
    "earlier events, and be manipulated by adversarial inputs.\n\n"
    "This project addresses this with a structured approach: handcrafted scenario schemas define the game world "
    "(rooms, NPCs, items, evidence chains) while the LLM provides creative, atmospheric narration within those "
    "constraints. The platform is built as a TypeScript monorepo with Next.js 15 + React 19 for the frontend "
    "and Express.js for the backend, using OpenAI GPT-5.4 for streaming narration. Six themed scenarios "
    "(noir detective, haunted manor, space station, pirate, western, cyberpunk) are available at launch. "
    "Planned features include multiplayer co-op via Socket.IO, AI-generated scene visuals, and deployment on "
    "Google Cloud Platform."
)

for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and p.text.strip() == "Introduction":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np, intro_text)
                break
        break

# ── 6. CONCEPT > OBJECTIVES ─────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Objectives":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "• Design and implement a structured scenario system where handcrafted JSON schemas "
                    "define game worlds (rooms, NPCs, items, evidence) while the LLM narrates freely.\n"
                    "• Build a full-stack TypeScript monorepo (Next.js frontend + Express backend + shared "
                    "types) with real-time streaming narration via SSE.\n"
                    "• Integrate OpenAI GPT-5.4 for atmospheric, context-aware narration and GPT-5-nano "
                    "for action suggestion generation.\n"
                    "• Create 6 distinct themed scenarios (Noir, Haunted Manor, Space Station, Pirate, "
                    "Western, Cyberpunk), each with unique rooms, NPCs, and evidence chains.\n"
                    "• Implement Socket.IO-based multiplayer co-op allowing 2-4 players to share a session.\n"
                    "• Add AI-generated scene visuals via gpt-image-1.5 for room transitions.\n"
                    "• Deploy to Google Cloud Run (backend) and Firebase Hosting (frontend) for public access."
                )
                break
        break

# ── 7. CONCEPT > BACKGROUND ─────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Background":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "Interactive fiction has been studied extensively. Montfort (2003) defines it as "
                    "'work in which a program accepts text input from a reader and prints text in response' [1]. "
                    "The genre peaked commercially in the 1980s with Infocom's parser-based games and was renewed "
                    "with Inform 7 and Twine [2].\n\n"
                    "The application of neural language models to interactive fiction was explored by Ammanabrolu "
                    "et al. (2020) [3], who showed transformer models could generate contextually relevant story "
                    "continuations. AI Dungeon (2019) was the first commercially successful LLM-driven text "
                    "adventure using GPT-2 and later GPT-3 [4], but its lack of game state led to well-documented "
                    "consistency failures.\n\n"
                    "OpenAI's GPT-5 family (2025) provides advanced streaming, function calling, and multi-modal "
                    "capabilities [5]. This project leverages GPT-5.4 for high-quality narration with streaming "
                    "output via Server-Sent Events, and GPT-5-nano for lightweight action suggestion generation. "
                    "The structured scenario approach ensures narrative consistency without sacrificing creative "
                    "freedom.\n\n"
                    "Procedural content generation in games has been studied by Shaker et al. (2016) [6]. While "
                    "this project uses handcrafted scenarios rather than fully procedural generation, the scenario "
                    "schema system is designed to support future LLM-based scenario generation."
                )
                break
        break

# ── 8. METHODOLOGY ──────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Methodology":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "The system is built as a TypeScript monorepo with three packages: @akboys/shared "
                    "(common types and scenario data), @akboys/server (Express.js backend), and @akboys/web "
                    "(Next.js 15 + React 19 frontend).\n\n"
                    "Scenario Design: Each of the 6 scenarios is defined as a typed TypeScript object containing "
                    "rooms (with exits, descriptions, investigation targets), NPCs (with dialogue trees, locations, "
                    "knowledge), and items (with evidence flags). The shared package exports a Scenario type that "
                    "enforces structural consistency across all scenarios.\n\n"
                    "AI Narration: The Express backend sends the full scenario context (minus the solution) as a "
                    "system prompt to OpenAI GPT-5.4. Player messages are appended to the conversation history "
                    "and streamed back via Server-Sent Events (SSE). The LLM narrates atmospherically while "
                    "respecting the scenario's world rules. A separate GPT-5-nano call generates 3 contextual "
                    "action suggestions after each narrator response.\n\n"
                    "Session Management: Each game session is identified by a UUID. The SessionStore interface "
                    "abstracts storage — currently in-memory (Map), planned migration to Firestore for "
                    "persistence across server restarts.\n\n"
                    "Multiplayer (Planned): Socket.IO will provide real-time bidirectional communication. "
                    "Players join a session by sharing its URL. The server broadcasts narrator responses to all "
                    "connected clients. Turn management ensures orderly gameplay.\n\n"
                    "Deployment (Planned): The backend will be containerized and deployed to Google Cloud Run "
                    "(serverless, auto-scaling). The frontend will be hosted on Firebase Hosting (global CDN). "
                    "Session data will persist in Firestore (NoSQL, same GCP project for minimal latency)."
                )
                break
        break

# ── 9. WORK PACKAGES ────────────────────────────────────────────
wps = [
    {
        "number": "WP1", "title": "Research, Setup & Demo (COMPLETED)",
        "start": "Week 1", "end": "Week 2",
        "participants": ["Batuhan Karaman", "Kadir Yigit Ozcelik", "Serdar Yengil", "Ata Berke Goktekin"],
        "weeks":        ["1-2", "1-2", "1-2", "1-2"],
        "objectives": "Set up GitHub repository, CI pipeline, and build initial demo.",
        "tasks": "T1.1 (W1) GitHub repo, branch protection, issue tracking\n"
                 "T1.2 (W1) Initial demo with Tkinter + Gemini API\n"
                 "T1.3 (W2) Architecture design (Structured Chaos concept)\n"
                 "T1.4 (W2) Project proposal and registration form",
        "deliverables": "D1.1 GitHub repo\nD1.2 Working demo (5 rooms, 3 NPCs)\nD1.3 Proposal document",
        "milestones": "M1.1 (W2) Demo ready — COMPLETED",
    },
    {
        "number": "WP2", "title": "Core Engine, Scenarios & Web Platform (COMPLETED)",
        "start": "Week 3", "end": "Week 6",
        "participants": ["Batuhan Karaman", "Kadir Yigit Ozcelik", "Serdar Yengil", "Ata Berke Goktekin"],
        "weeks":        ["3-6", "3-6", "3-6", "3-6"],
        "objectives": "Migrate from Python demo to TypeScript monorepo. Build full web platform with 6 scenarios.",
        "tasks": "T2.1 (W3) TypeScript monorepo setup (shared, server, web packages)\n"
                 "T2.2 (W3-4) Express.js backend with REST API (6 endpoints)\n"
                 "T2.3 (W4) Next.js 15 + React 19 frontend with noir-themed UI\n"
                 "T2.4 (W4-5) OpenAI GPT-5.4 streaming narration via SSE\n"
                 "T2.5 (W5) GPT-5-nano action suggestion system\n"
                 "T2.6 (W5-6) 6 scenario scripts (Noir, Haunted, Space, Pirate, Western, Cyberpunk)\n"
                 "T2.7 (W6) Session management with UUID routing, markdown rendering",
        "deliverables": "D2.1 TypeScript monorepo\nD2.2 6 playable scenarios\n"
                        "D2.3 Full-stack web app with streaming AI",
        "milestones": "M2.1 (W4) Backend API serving sessions\nM2.2 (W6) Web app fully playable — COMPLETED",
    },
    {
        "number": "WP3", "title": "Multiplayer & Persistence",
        "start": "Week 7", "end": "Week 10",
        "participants": ["Kadir Yigit Ozcelik", "Serdar Yengil", "Batuhan Karaman", ""],
        "weeks":        ["7-10", "7-10", "7-8", ""],
        "objectives": "Add real-time multiplayer via Socket.IO, migrate to Firestore, implement GameState tracking.",
        "tasks": "T3.1 (W7) Socket.IO server setup and room-based sessions\n"
                 "T3.2 (W7-8) Multiplayer UI (player list, typing indicators, message attribution)\n"
                 "T3.3 (W8) Explicit GameState tracking on server (room, inventory, visited)\n"
                 "T3.4 (W8-9) Migrate SessionStore from in-memory Map to Firestore\n"
                 "T3.5 (W9-10) Turn management and multiplayer session synchronization\n"
                 "T3.6 (W10) Firebase Auth for user authentication",
        "deliverables": "D3.1 Socket.IO multiplayer backend\nD3.2 Multiplayer frontend UI\n"
                        "D3.3 Firestore persistence\nD3.4 Firebase Auth integration",
        "milestones": "M3.1 (W8) Two players in same session\nM3.2 (W10) Full multiplayer co-op with auth",
    },
    {
        "number": "WP4", "title": "Advanced Features & Deployment",
        "start": "Week 9", "end": "Week 12",
        "participants": ["Ata Berke Goktekin", "Batuhan Karaman", "", ""],
        "weeks":        ["9-12", "9-12", "", ""],
        "objectives": "Deploy to GCP, add AI image generation, implement game-over conditions.",
        "tasks": "T4.1 (W9) Dockerfile for backend, Cloud Run deployment\n"
                 "T4.2 (W9-10) Firebase Hosting for frontend\n"
                 "T4.3 (W10-11) gpt-image-1.5 scene visualization with per-room cache\n"
                 "T4.4 (W11) Game-over conditions (win: solve mystery, lose: turn limit)\n"
                 "T4.5 (W11-12) Accusation mechanic and evidence tracking\n"
                 "T4.6 (W12) CI/CD pipeline for automatic deployments",
        "deliverables": "D4.1 Deployed web app (public URL)\nD4.2 AI image generation pipeline\n"
                        "D4.3 Game-over and win/lose system",
        "milestones": "M4.1 (W10) App publicly accessible\nM4.2 (W12) All advanced features integrated",
    },
    {
        "number": "WP5", "title": "Testing, Polish & Documentation",
        "start": "Week 13", "end": "Week 15",
        "participants": ["Batuhan Karaman", "Kadir Yigit Ozcelik", "Serdar Yengil", "Ata Berke Goktekin"],
        "weeks":        ["13-15", "13-15", "13-15", "13-15"],
        "objectives": "End-to-end testing, UX polish, poster, final report, and demo preparation.",
        "tasks": "T5.1 (W13) End-to-end playtesting across all 6 scenarios\n"
                 "T5.2 (W13) Multiplayer stress testing (concurrent sessions)\n"
                 "T5.3 (W14) UI/UX polish, performance tuning\n"
                 "T5.4 (W14) Poster design\n"
                 "T5.5 (W15) Final report and demo video",
        "deliverables": "D5.1 Test report\nD5.2 Poster\nD5.3 Final report\nD5.4 Demo video",
        "milestones": "M5.1 (W14) Poster submitted\nM5.2 (W15) Final demo ready",
    },
]

for ti, table in enumerate(doc.tables):
    if "Work package number" in table.rows[0].cells[0].text:
        wp_table = table
        wp_table_idx = ti
        break

def fill_wp_table(table, wp):
    def sc(ri, ci, text, bold=False):
        seen = set()
        unique_cells = []
        for c in table.rows[ri].cells:
            cid = id(c._tc)
            if cid not in seen:
                seen.add(cid)
                unique_cells.append(c)
        if ci < len(unique_cells):
            set_cell(unique_cells[ci], text, bold=bold)

    sc(0, 1, wp["number"], bold=True)
    sc(0, 3, f"{wp['start']} – {wp['end']}")
    sc(1, 1, wp["title"], bold=True)

    for col_i, name in enumerate(wp["participants"]):
        if col_i + 1 < len(table.rows[3].cells):
            set_cell(table.rows[3].cells[col_i + 1], name)
    for col_i, weeks in enumerate(wp["weeks"]):
        if col_i + 1 < len(table.rows[4].cells):
            set_cell(table.rows[4].cells[col_i + 1], weeks)

    sc(5, 0, "Objectives\n" + wp["objectives"])
    sc(6, 0, "Description of work\n" + wp["tasks"])
    sc(7, 0, "Deliverables\n" + wp["deliverables"])
    sc(8, 0, "Milestones\n" + wp["milestones"])

fill_wp_table(wp_table, wps[0])

from lxml import etree
for wp in wps[1:]:
    new_tbl = copy.deepcopy(wp_table._tbl)
    wp_table._tbl.addnext(new_tbl)
    new_table = None
    for t in doc.tables:
        if t._tbl is new_tbl:
            new_table = t
            break
    if new_table:
        fill_wp_table(new_table, wp)
        wp_table = new_table

# ── 10. DEMONSTRATION ───────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Demonstration" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "The final demonstration will show a live playthrough in a web browser in two modes: "
                    "single-player and multiplayer co-op. The single-player demo will cover at least two "
                    "of the six scenarios (e.g., Noir and Cyberpunk) from scenario selection to mystery "
                    "resolution. The multiplayer demo will have two participants collaborating in the same "
                    "session via shared URL.\n\n"
                    "Performance measures: (1) Streaming latency — narrator response begins within 1 second. "
                    "(2) Scenario consistency — NPCs, rooms, and items behave according to their definitions. "
                    "(3) Multiplayer synchronization — all players see the same game state in real time. "
                    "(4) Visual coherence — generated scene images match the scenario theme. "
                    "(5) Deployment — the app is publicly accessible via a URL without local setup."
                )
                break
        break

# ── 11. IMPACT ──────────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Impact" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "This project advances AI-powered interactive entertainment by combining structured "
                    "scenario design with creative LLM narration. The TypeScript monorepo architecture "
                    "demonstrates a professional, full-stack approach to AI application development.\n\n"
                    "The platform can serve as an educational tool for language learning, creative writing, "
                    "and critical thinking through mystery-solving gameplay. The multiplayer co-op mode "
                    "develops teamwork and collaborative problem-solving. The scenario schema system is "
                    "extensible — new themes and stories can be added by defining new scenario objects "
                    "without modifying the engine. The project will be open-sourced on GitHub."
                )
                break
        break

# ── 12. RISK ANALYSIS ───────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Risk" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "Risk 1 — LLM narration inconsistency (Likelihood: Medium, Impact: Medium): "
                    "The LLM may contradict earlier events or hallucinate items. Mitigated by including "
                    "full conversation history in each API call and scenario context in the system prompt. "
                    "Planned server-side GameState tracking will add a second layer of validation.\n\n"
                    "Risk 2 — OpenAI API costs (Likelihood: Medium, Impact: Medium): "
                    "GPT-5.4 streaming + GPT-5-nano suggestions + gpt-image-1.5 could accumulate costs. "
                    "Mitigated by using nano model for suggestions (low cost), caching images per room, "
                    "and setting per-session token limits.\n\n"
                    "Risk 3 — Socket.IO multiplayer latency (Likelihood: Low, Impact: Medium): "
                    "Real-time synchronization may lag under load. Mitigated by turn-based gameplay "
                    "(no real-time twitch action required) and Cloud Run auto-scaling.\n\n"
                    "Risk 4 — Data loss with in-memory store (Likelihood: High, Impact: High): "
                    "Current in-memory SessionStore loses all data on restart. ACTIVE risk — Firestore "
                    "migration (Issue #2) is the top priority for WP3.\n\n"
                    "Risk 5 — Scope creep (Likelihood: Medium, Impact: Medium): "
                    "WP1-WP2 already deliver a complete, playable web app. WP3-WP5 add enhancements. "
                    "If multiplayer proves too complex, the single-player app is a viable final product."
                )
                break
        break

# ── 13. GANTT CHART ─────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Gantt" in p.text:
        for j in range(i+1, min(i+6, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                for run in np.runs:
                    run.text = ""
                if not np.runs:
                    np.add_run()
                np.runs[0].text = ""
                r = np.add_run()
                r.add_picture(_gantt_path, width=Inches(5.8))
                np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if j+1 < len(doc.paragraphs):
                    set_para(doc.paragraphs[j+1], "")
                break
        break

# ── 14. ECONOMICAL & ETHICAL ─────────────────────────────────────
econ_texts = [
    "The project uses OpenAI API for LLM narration (GPT-5.4), action suggestions (GPT-5-nano), and "
    "planned image generation (gpt-image-1.5). OpenAI API costs are usage-based; development costs are "
    "managed through the team's shared API budget. Deployment uses Google Cloud Platform: Cloud Run for "
    "the backend (serverless, scales to zero when idle — zero cost at rest), Firebase Hosting for the "
    "frontend (free tier), and Firestore for the database (free tier: 50K reads/day, 20K writes/day). "
    "No special hardware is required beyond standard development laptops. The domain and hosting "
    "infrastructure are covered by GCP free tier and Firebase free plan.",

    "Content Safety: The scenario schema defines all available rooms, NPCs, and items. The LLM narrates "
    "within these boundaries. Content moderation can be added as a further safeguard via OpenAI's "
    "moderation API.\n\n"
    "Intellectual Property: All scenarios are original creations by the team. AI-generated narration "
    "and images are clearly disclosed as AI-generated.\n\n"
    "Player Wellbeing: The platform includes session identifiers but does not employ dark patterns "
    "such as artificial urgency or manipulative reward schedules.\n\n"
    "Data Privacy: Firebase Auth handles user identity securely. Conversation logs are stored in "
    "Firestore within the team's GCP project. No data is shared with third parties beyond OpenAI's "
    "API processing (governed by OpenAI's data usage policy)."
]

for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and "Economical" in p.text:
        count = 0
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                if count < len(econ_texts):
                    set_para(np, econ_texts[count])
                    count += 1
        break

# ── 15. REFERENCES ───────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and "References" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "[1] N. Montfort, Twisty Little Passages: An Approach to Interactive Fiction. MIT Press, 2003.\n"
                    "[2] A. Plotkin, The Inform Designer's Handbook, 2004. Available: http://inform-fiction.org\n"
                    "[3] P. Ammanabrolu et al., How to Avoid Being Eaten by a Grue, arXiv:2006.07409, 2020.\n"
                    "[4] N. Walton, AI Dungeon: Play any adventure you can imagine. Latitude, 2019. Available: https://aidungeon.io\n"
                    "[5] OpenAI, GPT-5 API Documentation, 2025. Available: https://platform.openai.com/docs\n"
                    "[6] N. Shaker, J. Togelius, M. Nelson, Procedural Content Generation in Games. Springer, 2016. Available: http://pcgbook.com\n"
                    "[7] Socket.IO Documentation, 2024. Available: https://socket.io/docs/v4/"
                )
                break
        break

# ── Save ─────────────────────────────────────────────────────────
out = "/Users/batuhankaraman/comp491-akboys/docs/filled/COMP491_Proposal_Updated_FILLED.docx"
doc.save(out)
print(f"Saved: {out}")
