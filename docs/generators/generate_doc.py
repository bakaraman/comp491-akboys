from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Style helpers
def set_run(run, size=10, bold=False, color=RGBColor(0,0,0), font_name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name

def add_heading_text(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p

def add_subheading(text, size=11):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=size, bold=True)
    return p

def add_body(text, size=10):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=size)
    p.paragraph_format.line_spacing = Pt(14)
    return p

def add_code(text, size=9):
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    set_run(run, size=size, color=RGBColor(40,40,40), font_name="Consolas")
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=9, bold=True)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run(run, size=9)
    doc.add_paragraph()  # spacing
    return table

# ===== TITLE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_after = Pt(2)
run = title.add_run("Text-Based Adventure with LLMs")
set_run(run, size=16, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.space_after = Pt(2)
run = sub.add_run("Technical Overview")
set_run(run, size=11, color=RGBColor(80,80,80))

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.space_after = Pt(2)
run = info.add_run("COMP 491 — Computer Engineering Design, Spring 2026")
set_run(run, size=9, color=RGBColor(100,100,100))

team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
team.space_after = Pt(2)
run = team.add_run("Kadir Yigit Ozcelik (79975) · Serdar Yengil (80232) · Batuhan Karaman (79791) · Ata Berke Goktekin (80277)")
set_run(run, size=9)

adv = doc.add_paragraph()
adv.alignment = WD_ALIGN_PARAGRAPH.CENTER
adv.space_after = Pt(8)
run = adv.add_run("Advisor: Baris Akgun")
set_run(run, size=9)

# Separator
sep = doc.add_paragraph()
sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sep.add_run("—" * 50)
set_run(run, size=8, color=RGBColor(180,180,180))

# ===== 1. WHAT WE BUILD =====
add_heading_text("1. Project Overview")
add_body(
    "This project is a platform for AI-powered text adventure games where players type natural language commands and an LLM narrates the story. "
    "The platform supports multiple game types and universes (noir detective, cyberpunk, fantasy, horror, and more). "
    "The LLM is only the storyteller. A Python engine acts as the referee and holds all game rules. "
    "The LLM can only affect the game world through validated function calls. "
    "The AI brings creative narration while Python guarantees the game never breaks. "
    "Every playthrough is unique because the AI generates a fresh scenario each time. "
    "Players can type free text, pick from suggested actions, or use voice input."
)

# ===== 2. HOW IT WORKS =====
add_heading_text("2. How It Works")
add_body(
    'The player types something like "examine the body". The LLM reads this and decides to call '
    'investigate(target="body"). Python checks: does this target exist in the current room? Does the player '
    "meet the prerequisites? If yes, Python returns a JSON result with the description and any new evidence. "
    "The LLM takes this JSON and turns it into atmospheric, literary narration. The player never sees raw data."
)
add_code("Player Input → LLM picks a function call if necessary → Python validates & updates state → LLM narrates → Player Output")

add_subheading("2.1 Function Calls (Demo Example)")
add_body(
    "The LLM can only change the game state through predefined function calls. "
    "The set of available functions depends on the game type and scenario. "
    "In the current demo (a 1920s noir detective game), these 6 functions serve as an example:"
)
add_table(
    ["Function", "What it does", "What Python validates"],
    [
        ["maps(direction)", "Move between rooms", "Does this exit exist?"],
        ["investigate(target)", "Examine an object or area", "Is target here? Prerequisite evidence?"],
        ["pickup_item(item)", "Put item into inventory", "Was it discovered via investigate?"],
        ["interact_npc(name, action)", "Talk to / question / threaten NPC", "Is NPC in the same room?"],
        ["accuse_suspect(who, evidence)", "Final accusation (ends game)", "Evidence in inventory? Correct combo?"],
        ["take_damage(amount, reason)", "Reduce player sanity", "Amount capped 1-50. Game over at 0."],
    ]
)

add_body(
    "A different game type would have different functions. For example, a cyberpunk scenario might replace "
    "accuse_suspect with hack_terminal and scan_biometrics. A fantasy game might add cast_spell. "
    "The core principle stays the same: every function validates before it allows anything."
)

add_subheading("2.2 Concrete Example: Evidence Chain (from the demo)")
add_body(
    "In the working demo, the player solves a 1920s noir murder. The solution requires steps in order. "
    "The LLM cannot skip any step because each function validates prerequisites:"
)
add_code(
    '1. investigate("body") in Study Room\n'
    '   → Python grants "poison_signs" evidence\n\n'
    '2. investigate("bushes") in Garden\n'
    '   → Python checks: does player have "poison_signs"?\n'
    '   → If no: "You need a clue first." If yes: reveals poison bottle\n\n'
    '3. pickup_item("empty_poison_bottle")\n'
    '   → Python checks: was this item discovered? If yes: added to inventory\n\n'
    '4. accuse_suspect("jenkins", "empty_poison_bottle")\n'
    '   → Correct suspect + correct evidence = YOU WIN'
)
add_body(
    'Even if the LLM hallucinates "the player found a gun", there is no add_to_inventory() function. '
    "The only path to inventory is pickup_item, which validates against discovered items."
)

# ===== 3. SANDBOX =====
add_heading_text("3. Three Layers of Sandbox")
add_body(
    'Layer 1 — System Prompt: The LLM receives a strict rule: "You are the narrator. Never change game state yourself. '
    'All mechanics run through function calls." But prompts can be jailbroken, so this layer alone is not enough.'
)
add_body(
    "Layer 2 — Function Call Gating: This is the real enforcement. The LLM physically cannot modify game state "
    "except through validated functions defined for that game type. Each function checks rules before it allows anything. "
    "Wrong room? Rejected. Missing prerequisite? Rejected. Item not discovered? Rejected."
)
add_body(
    "Layer 3 — Scenario Schema: The entire game world is a structured JSON. Rooms, exits, NPCs, items, "
    "evidence chains, and the solution are all pre-defined. The Python referee only allows actions within this schema. "
    "The LLM has freedom in HOW it describes events, never in WHAT events happen."
)

# ===== 4. PROCEDURAL =====
add_heading_text("4. Procedural Story Generation")
add_body(
    "The system uses the LLM in two separate phases. In Phase 1 (before the game), the LLM acts as a scenario architect. "
    'The prompt is: "Generate a cyberpunk murder mystery with 6 rooms, 3 suspects, 2 key evidence items. Output as JSON." '
    "Python validates the JSON: are all rooms connected? Is the evidence chain solvable? Do NPC locations match? "
    "If validation fails, the system re-prompts with specific errors."
)
add_body(
    "In Phase 2 (during the game), a separate LLM session receives the scenario but NOT the solution. "
    "It narrates based on function call results, just like the player discovers the truth step by step. "
    "The narrator never knows who did it."
)
add_code(
    '{\n'
    '  "rooms": {"study": {"exits": {"south": "kitchen"}, "targets": {...}}},\n'
    '  "npcs":  {"jenkins": {"location": "kitchen", "is_guilty": true}},\n'
    '  "items": {"poison_bottle": {"is_key_evidence": true}},\n'
    '  "solution": {"guilty": "jenkins", "evidence": "poison_bottle"}\n'
    '}'
)
add_body(
    "Each game template (noir, cyberpunk, fantasy, horror) provides a style prefix for the narrator, "
    "a generation prompt for the scenario architect, mechanic modifiers (horror makes sanity drop faster), "
    "and extra function calls (cyberpunk adds hack_terminal, fantasy adds cast_spell)."
)

# ===== 5. VISUALS =====
add_heading_text("5. AI Visual Generation")
add_body(
    "Every room transition triggers an image generation request. The prompt is composed from a template style prefix "
    'plus the room description. For example, the noir template prepends "1920s noir ink illustration, dark shadows, '
    'sepia tones:" to every room description. This keeps all images in the same visual style throughout the game. '
    "Images are cached per room (no re-generation on revisit), generated asynchronously (the game never freezes), "
    "and have a fallback placeholder if generation fails. The image model is Gemini 2.5 Flash (native image generation)."
)

# ===== 6. MULTIPLAYER =====
add_heading_text("6. Multiplayer: Queue-Based Co-op")
add_body(
    "Three approaches were evaluated: strict turn-based (too slow), free-for-all (race conditions, LLM confusion), "
    "and queue-based. The chosen approach is queue-based: players submit actions anytime, actions enter a FIFO queue, "
    "Python processes them one by one, and the narration is broadcast to all players via WebSocket."
)
add_code(
    'Player A: "examine body"     → Queue pos 1 → Process → Broadcast\n'
    'Player B: "talk to Jenkins"  → Queue pos 2 → Process → Broadcast\n'
    'Player A: "go south"         → Queue pos 3 → Process → Broadcast'
)
add_body(
    "The map, evidence board, NPCs, and story progress are shared. Each player has their own location, inventory, sanity, and role."
)
add_table(
    ["Role", "Special Ability", "Limitation"],
    [
        ["Detective", "Can accuse suspects, dust for prints", "None"],
        ["Journalist", "Bonus NPC dialogue options", "Cannot accuse"],
        ["Doctor", "Can analyze substances, partial sanity immunity", "Cannot threaten NPCs"],
        ["Thief", "Can pick locks, access hidden areas", "Loses extra sanity"],
    ]
)
add_body(
    "Only the Detective can make the final accusation. NPCs remember interactions with all players: "
    "if Player A threatens Jenkins, Jenkins is hostile to Player B too."
)

# ===== 7. TECH & TIMELINE =====
add_heading_text("7. Tech Stack & Timeline")
add_table(
    ["Layer", "Technology"],
    [
        ["LLM & Images", "Vertex AI — Gemini 2.5 Flash (function calls + native image generation)"],
        ["Backend", "Google Cloud Run (FastAPI, serverless)"],
        ["Frontend", "React + Next.js (Vercel)"],
        ["Multiplayer", "WebSocket (Socket.IO)"],
        ["Database", "PostgreSQL (Cloud SQL)"],
        ["Cloud Credits", "GCP $300 free credit — 3 months"],
    ]
)
add_table(
    ["Weeks", "Milestone"],
    [
        ["1-2", "Project setup, proposal, GitHub repo"],
        ["3-4", "Procedural story generation, schema validation"],
        ["5-6", "Web UI (React + FastAPI), basic gameplay"],
        ["7-8", "AI visual generation, cache system"],
        ["9-10", "Multiplayer backend (WebSocket, queue)"],
        ["11-12", "Multiplayer frontend, roles, lobby"],
        ["13-14", "Polish, adversarial tests, poster"],
        ["15", "Poster event, final report, demo"],
    ]
)

# Save
out = "/Users/batuhankaraman/demo491/project_technical_document.docx"
doc.save(out)
print(f"Saved to {out}")
