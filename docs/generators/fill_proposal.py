import os
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Generate Gantt chart PNG ─────────────────────────────────────
_gantt_path = "/tmp/gantt_chart.png"

def _make_gantt():
    fig, ax = plt.subplots(figsize=(12, 5))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("#f8f8f8")

    tasks = [
        ("WP1: Research & Setup",           1,  2, "#4e79a7"),
        ("WP2: Story Generation & Engine",  3,  6, "#f28e2b"),
        ("WP3: Frontend & Visuals",         5,  8, "#59a14f"),
        ("WP4: Multiplayer System",         9, 12, "#e15759"),
        ("WP5: Testing & Polish",          13, 15, "#76b7b2"),
    ]
    milestones = [
        ("M1.1 Demo Ready",     2,  0),
        ("M2.2 FastAPI Live",   6,  1),
        ("M3.2 Visuals Live",   8,  2),
        ("M4.2 4-Player Co-op",12,  3),
        ("M5.1 Poster",        14,  4),
        ("M5.2 Final Demo",    15,  4),
    ]
    y_pos = list(range(len(tasks) - 1, -1, -1))

    for i, (label, start, end, color) in enumerate(tasks):
        y = y_pos[i]
        ax.barh(y, end - start + 1, left=start - 0.5, height=0.55,
                color=color, alpha=0.85, edgecolor="white", linewidth=1.2)
        ax.text(start - 0.5 + (end - start + 1) / 2, y,
                f"W{start}–W{end}", ha="center", va="center",
                fontsize=8, color="white", fontweight="bold")

    for label, week, task_i in milestones:
        y = y_pos[task_i]
        ax.plot(week, y, marker="D", color="#333333", markersize=7, zorder=5)
        ax.text(week, y + 0.33, label, ha="center", va="bottom",
                fontsize=6.5, color="#333333", fontstyle="italic")

    ax.set_xlim(0.5, 15.5)
    ax.set_xticks(range(1, 16))
    ax.set_xticklabels([f"W{i}" for i in range(1, 16)], fontsize=8)
    ax.set_yticks(y_pos)
    ax.set_yticklabels([t[0] for t in tasks], fontsize=9)
    ax.xaxis.grid(True, linestyle="--", alpha=0.5, color="#cccccc")
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)
    ax.axvline(x=15, color="#e15759", linestyle=":", linewidth=1.5, alpha=0.6)
    ax.axvline(x=13, color="#76b7b2", linestyle=":", linewidth=1.5, alpha=0.6)
    ax.set_xlabel("Week", fontsize=9, labelpad=6)
    ax.set_title("Project Schedule — COMP 491 Spring 2026",
                 fontsize=11, fontweight="bold", pad=10)
    plt.tight_layout(pad=1.2)
    plt.savefig(_gantt_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close()

_make_gantt()

doc = Document("/Users/batuhankaraman/demo491/lecture_slides/COMP491_Proposal_Template.docx")

# ── helpers ─────────────────────────────────────────────────────
def set_para(para, text, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Clear all runs in a paragraph and set fresh text."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run()
    r.text = text
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    para.alignment = align

def set_cell(cell, text, size=10, bold=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""
    p = cell.paragraphs[0]
    if p.runs:
        r = p.runs[0]
    else:
        r = p.add_run()
    r.text = text
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "Times New Roman"

def add_row_to_table(table, values, size=10):
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            set_cell(row.cells[i], val, size=size)
    return row

def insert_para_after(para, text, style_name=None, size=12, bold=False,
                       italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       space_before=0, space_after=6):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    para._p.addnext(new_p)
    # Find the newly created paragraph object
    idx = list(doc.paragraphs).index(para) + 1
    new_para = doc.paragraphs[idx]
    if style_name:
        try:
            new_para.style = doc.styles[style_name]
        except Exception:
            pass
    new_para.paragraph_format.space_before = Pt(space_before)
    new_para.paragraph_format.space_after = Pt(space_after)
    new_para.alignment = align
    r = new_para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0, 0, 0)
    return new_para

# ── 1. TITLE PAGE ────────────────────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip() == "PROJECT TITLE":
        set_para(p, "Text-Based Adventure with LLMs",
                 size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    elif p.text.strip() == "Fall 2025":
        set_para(p, "Spring 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    elif p.text.strip() == "Project Advisor:":
        set_para(p, "Project Advisor: Baris Akgun")

# ── 2. PARTICIPANT TABLE ─────────────────────────────────────────
members = [
    ("Kadir Yigit Ozcelik", "79975", "kadirozcelik21@ku.edu.tr", "+90 551 215 64 74"),
    ("Serdar Yengil",       "80232", "syengil21@ku.edu.tr",      "+90 545 775 08 63"),
    ("Batuhan Karaman",     "79791", "bkaraman21@ku.edu.tr",     "+90 534 878 90 16"),
    ("Ata Berke Goktekin",  "80277", "agoktekin21@ku.edu.tr",    "+90 538 064 49 44"),
]
for table in doc.tables:
    if table.rows[0].cells[0].text.strip() == "Name":
        # Fill existing empty rows (rows 1-4) directly
        for row_i, m in enumerate(members):
            if row_i + 1 < len(table.rows):
                for col_i, val in enumerate(m):
                    set_cell(table.rows[row_i + 1].cells[col_i], val)
        break

# ── 3. ABSTRACT ──────────────────────────────────────────────────
for p in doc.paragraphs:
    if "A brief summary of your project." in p.text:
        set_para(p, "")
    elif "Summarize the motivation" in p.text:
        set_para(p,
            "This project is a platform for AI-powered text adventure games. "
            "Players explore procedurally generated story worlds using natural language. "
            "A Large Language Model (LLM) acts as the narrator while a Python engine enforces all game rules "
            "through a validated function call interface. The key innovation is the \"Structured Chaos\" "
            "architecture: the LLM has full creative freedom in narration but cannot alter game state outside "
            "of validated function calls, solving the consistency and sandbox problems found in systems like "
            "AI Dungeon. The platform supports multiple universe templates, AI-generated scene visuals using "
            "Gemini 2.5 Flash, and a queue-based multiplayer co-op mode where up to 4 players assume "
            "different roles in the same story."
        )

# ── 4. TABLE OF CONTENTS ─────────────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip() == "Table of contents":
        set_para(p, "Table of Contents", bold=True, size=12,
                 align=WD_ALIGN_PARAGRAPH.LEFT)

# ── helper: find paragraph by text fragment ──────────────────────
def find_para(fragment):
    for p in doc.paragraphs:
        if fragment in p.text:
            return p
    return None

# ── 5. INTRODUCTION ──────────────────────────────────────────────
p_intro_placeholder = find_para("background information and literature")
if not p_intro_placeholder:
    # Try the next normal paragraph after Introduction heading
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "mh1" and "Introduction" in p.text:
            # next normal para
            for j in range(i+1, len(doc.paragraphs)):
                if doc.paragraphs[j].text.strip():
                    p_intro_placeholder = doc.paragraphs[j]
                    break
            break

intro_text = (
    "Text-based adventure games, also known as interactive fiction, are one of the oldest forms of digital "
    "entertainment. Pioneered by Colossal Cave Adventure (1976) and Zork (1977), these games let players explore "
    "narratives through natural language. Despite their long history, the genre has been limited by hand-authored "
    "content: every room, NPC, and puzzle had to be written by a developer in advance.\n\n"
    "The emergence of Large Language Models (LLMs) opens a new chapter for this genre. AI Dungeon (2019) "
    "demonstrated that LLMs can generate coherent narratives on demand, but purely LLM-driven games suffer from "
    "a fundamental problem: the model has no concept of rules or game state. It can hallucinate items, contradict "
    "earlier events, and be manipulated by adversarial inputs.\n\n"
    "This project addresses this with the \"Structured Chaos\" architecture: a Python engine enforces all rules "
    "while the LLM narrates. The platform supports multiple universes (noir detective, cyberpunk, fantasy, horror), "
    "procedurally generates a unique scenario every session, renders AI visuals per scene via Gemini 2.5 Flash, "
    "and supports multiplayer co-op with up to 4 players."
)

# Find Introduction heading, then replace following Normal paragraph
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and p.text.strip() == "Introduction":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np, intro_text)
                break
        break

# ── 6. CONCEPT > OBJECTIVES ─────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Objectives":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "• Design and implement the Structured Chaos architecture separating the LLM narrator "
                    "from the Python referee engine.\n"
                    "• Build a procedural scenario generation system where the LLM creates a unique, "
                    "validated game world before each session.\n"
                    "• Develop a three-layer sandbox enforcement mechanism preventing the LLM from "
                    "breaking game rules.\n"
                    "• Integrate Gemini 2.5 Flash native image generation to produce scene-appropriate "
                    "visuals for each room transition.\n"
                    "• Implement a queue-based multiplayer co-op system supporting up to 4 players with "
                    "distinct roles.\n"
                    "• Support at least 5 distinct universe templates (noir, cyberpunk, fantasy, horror, "
                    "western), each with its own tone and function call extensions."
                )
                break
        break

# ── 7. CONCEPT > BACKGROUND ─────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Background":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "Interactive fiction has been studied extensively. Montfort (2003) defines it as "
                    "'work in which a program accepts text input from a reader and prints text in response' [1]. "
                    "The genre peaked commercially in the 1980s with Infocom's parser-based games and was renewed "
                    "with Inform 7 and Twine [2].\n\n"
                    "The application of neural language models to interactive fiction was explored by Ammanabrolu "
                    "et al. (2020) [3], who showed transformer models could generate contextually relevant story "
                    "continuations. AI Dungeon (2019) was the first commercially successful LLM-driven text "
                    "adventure using GPT-2 and later GPT-3 [4], but its lack of game state led to well-documented "
                    "consistency failures.\n\n"
                    "Function calling in LLMs, introduced by OpenAI (2023) and supported by Google Gemini, "
                    "provides a structured mechanism for models to interact with external systems [5]. This project "
                    "leverages Gemini's function calling to create a strict, validated interface between the LLM "
                    "narrator and the Python referee, solving the consistency and sandbox problems of prior systems.\n\n"
                    "Procedural content generation in games has been studied by Shaker et al. (2016) [6]. This "
                    "project proposes an LLM-based PCG approach where the model generates structured JSON game "
                    "scenarios validated by the Python engine before gameplay begins."
                )
                break
        break

# ── 8. METHODOLOGY ──────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and p.text.strip() == "Methodology":
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "The system has two layers. The LLM Layer (Narrator) reads player input, selects a "
                    "function call if necessary, receives the Python result as JSON, and transforms it into "
                    "atmospheric narration. The Python Layer (Referee) holds the complete GameState and validates "
                    "every action before executing it.\n\n"
                    "The sandbox is enforced through three independent layers:\n"
                    "Layer 1 — System Prompt: The LLM receives a rule not to modify state directly. "
                    "This layer alone is insufficient since prompts can be jailbroken.\n"
                    "Layer 2 — Function Call Gating: The LLM can only affect game state through predefined, "
                    "validated Python functions. Wrong room? Rejected. Missing prerequisite? Rejected. "
                    "Item not discovered? Rejected.\n"
                    "Layer 3 — Scenario Schema: The entire game world is a JSON schema (rooms, exits, NPCs, "
                    "items, evidence chains, solution) defined before the game starts. The referee only permits "
                    "actions within this schema.\n\n"
                    "Procedural generation uses two LLM phases. Phase 1 (pre-game): the LLM generates the full "
                    "scenario JSON for the chosen template. Python validates it (graph connectivity, evidence "
                    "chain solvability, NPC locations). Phase 2 (gameplay): a separate LLM session narrates "
                    "based on function call results. It never receives the solution directly.\n\n"
                    "Multiplayer uses a queue-based approach: actions enter a FIFO queue, Python processes them "
                    "one by one, and narration is broadcast to all clients via WebSocket. Shared state (map, "
                    "evidence, NPCs) is common to all players; individual state (location, inventory, sanity, "
                    "role) is per-player. Roles: Detective (final accusation), Journalist (bonus NPC dialogue), "
                    "Doctor (substance analysis, sanity immunity), Thief (lock-picking, hidden areas).\n\n"
                    "Visual generation triggers on every room transition. A template-specific style prefix is "
                    "prepended to the room description and sent to Gemini 2.5 Flash's native image generation. "
                    "Images are cached per room and generated asynchronously."
                )
                break
        break

# ── 9. WORK PACKAGES ────────────────────────────────────────────
wps = [
    {
        "number": "WP1", "title": "Research, Setup & Demo Refinement",
        "start": "Week 1", "end": "Week 2",
        "participants": ["Batuhan Karaman", "Kadir Yigit Ozcelik", "Serdar Yengil", "Ata Berke Goktekin"],
        "weeks":        ["1-2", "1-2", "1-2", "1-2"],
        "objectives": "Set up GitHub repository, CI pipeline, and refine the existing working demo into a clean, documented codebase.",
        "tasks": "T1.1 (W1) GitHub repo, branch protection, issue tracking\nT1.2 (W1) Refactor demo: separate engine from LLM layer\nT1.3 (W2) Unit tests for all function calls\nT1.4 (W2) Architecture documentation",
        "deliverables": "D1.1 GitHub repo with CI\nD1.2 Refactored demo with tests\nD1.3 Architecture document",
        "milestones": "M1.1 (W2) Clean, tested demo ready for extension",
    },
    {
        "number": "WP2", "title": "Procedural Story Generation & Core Engine",
        "start": "Week 3", "end": "Week 6",
        "participants": ["Batuhan Karaman", "Ata Berke Goktekin", "", ""],
        "weeks":        ["3-6", "3-6", "", ""],
        "objectives": "Build the two-phase LLM story generation system and migrate to FastAPI backend.",
        "tasks": "T2.1 (W3) Design scenario JSON schema and validation rules\nT2.2 (W3) Implement Phase 1 LLM scenario generation with re-prompt loop\nT2.3 (W4) Build schema validator (graph connectivity, evidence chain)\nT2.4 (W4-5) Implement 5 universe templates\nT2.5 (W5-6) Migrate to FastAPI backend",
        "deliverables": "D2.1 Scenario generation system\nD2.2 5 universe templates\nD2.3 FastAPI backend",
        "milestones": "M2.1 (W4) LLM generates valid scenarios\nM2.2 (W6) FastAPI serving game sessions",
    },
    {
        "number": "WP3", "title": "Frontend & AI Visual Generation",
        "start": "Week 5", "end": "Week 8",
        "participants": ["Kadir Yigit Ozcelik", "Ata Berke Goktekin", "", ""],
        "weeks":        ["5-8", "5-8", "", ""],
        "objectives": "Build the React/Next.js frontend and integrate Gemini 2.5 Flash image generation.",
        "tasks": "T3.1 (W5-6) React + Next.js project setup, game UI components\nT3.2 (W6) Connect frontend to FastAPI backend\nT3.3 (W7) Integrate Gemini 2.5 Flash image generation\nT3.4 (W7) Async image generation with room-level cache\nT3.5 (W8) Style prefix system, fallback placeholder",
        "deliverables": "D3.1 React frontend\nD3.2 AI visual generation pipeline\nD3.3 Template style system",
        "milestones": "M3.1 (W6) Playable in browser\nM3.2 (W8) Scene visuals per room",
    },
    {
        "number": "WP4", "title": "Multiplayer System",
        "start": "Week 9", "end": "Week 12",
        "participants": ["Serdar Yengil", "Batuhan Karaman", "", ""],
        "weeks":        ["9-12", "9-12", "", ""],
        "objectives": "Implement the queue-based multiplayer co-op system with WebSocket and player roles.",
        "tasks": "T4.1 (W9) WebSocket server (Socket.IO)\nT4.2 (W9-10) Action queue system (FIFO, broadcast)\nT4.3 (W10) Shared vs individual state split\nT4.4 (W11) Player role system\nT4.5 (W11-12) Lobby, role selection, session management UI\nT4.6 (W12) NPC cross-player memory",
        "deliverables": "D4.1 WebSocket multiplayer backend\nD4.2 Role system\nD4.3 Multiplayer frontend",
        "milestones": "M4.1 (W10) Two players in same session\nM4.2 (W12) Full 4-player co-op with roles",
    },
    {
        "number": "WP5", "title": "Testing, Polish & Documentation",
        "start": "Week 13", "end": "Week 15",
        "participants": ["Batuhan Karaman", "Kadir Yigit Ozcelik", "Serdar Yengil", "Ata Berke Goktekin"],
        "weeks":        ["13-15", "13-15", "13-15", "13-15"],
        "objectives": "Adversarial testing of sandbox, UX polish, poster, and final report.",
        "tasks": "T5.1 (W13) Adversarial sandbox testing\nT5.2 (W13) End-to-end playtesting across all 5 templates\nT5.3 (W14) UI/UX polish, performance tuning\nT5.4 (W14) Poster design\nT5.5 (W15) Final report and demo preparation",
        "deliverables": "D5.1 Test report\nD5.2 Poster\nD5.3 Final report\nD5.4 Demo video",
        "milestones": "M5.1 (W14) Poster submitted\nM5.2 (W15) Final demo ready",
    },
]

# Find WP table and replace it completely
for ti, table in enumerate(doc.tables):
    if "Work package number" in table.rows[0].cells[0].text:
        wp_table = table
        wp_table_idx = ti
        break

def fill_wp_table(table, wp):
    def sc(ri, ci, text, bold=False):
        # Use first unique cell in this row (handles merged cells)
        seen = set()
        unique_cells = []
        for c in table.rows[ri].cells:
            cid = id(c._tc)
            if cid not in seen:
                seen.add(cid)
                unique_cells.append(c)
        if ci < len(unique_cells):
            set_cell(unique_cells[ci], text, bold=bold)

    # Row 0: WP number (unique col 1) and date range (unique col 3 = merged cols 4-5)
    sc(0, 1, wp["number"], bold=True)
    sc(0, 3, f"{wp['start']} – {wp['end']}")

    # Row 1: title goes in unique col 1 (merged cols 1-5)
    sc(1, 1, wp["title"], bold=True)

    # Rows 3-4: participant names and weeks (cols 1-5 are separate cells)
    for col_i, name in enumerate(wp["participants"]):
        if col_i + 1 < len(table.rows[3].cells):
            set_cell(table.rows[3].cells[col_i + 1], name)
    for col_i, weeks in enumerate(wp["weeks"]):
        if col_i + 1 < len(table.rows[4].cells):
            set_cell(table.rows[4].cells[col_i + 1], weeks)

    # Rows 5-8: entire row is ONE merged cell — write label + content together
    sc(5, 0, "Objectives\n" + wp["objectives"])
    sc(6, 0, "Description of work\n" + wp["tasks"])
    sc(7, 0, "Deliverables\n" + wp["deliverables"])
    sc(8, 0, "Milestones\n" + wp["milestones"])

# Fill first WP into existing table
fill_wp_table(wp_table, wps[0])

# For WP 2-5, copy the table XML and insert after
from lxml import etree
for wp in wps[1:]:
    # deep copy the wp table
    new_tbl = copy.deepcopy(wp_table._tbl)
    wp_table._tbl.addnext(new_tbl)
    # find the new table object
    new_table = None
    for t in doc.tables:
        if t._tbl is new_tbl:
            new_table = t
            break
    if new_table:
        fill_wp_table(new_table, wp)
        # Update wp_table reference so each copy is after the last
        wp_table = new_table

# ── 10. DEMONSTRATION ───────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Demonstration" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "The final demonstration will show a live playthrough in a web browser in two modes: "
                    "single-player and multiplayer co-op. The single-player demo will cover at least two "
                    "universe templates from scenario generation to final resolution. The multiplayer demo "
                    "will have two or more participants with different roles solving a mystery together.\n\n"
                    "Performance measures: (1) Sandbox integrity — no adversarial input can modify game state "
                    "outside validated function calls. (2) Scenario validity — 100% of generated scenarios "
                    "pass Python validation within two attempts. (3) Narrative coherence — each session "
                    "produces a distinct, internally consistent story. (4) Visual consistency — all generated "
                    "images match the active template style."
                )
                break
        break

# ── 11. IMPACT ──────────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Impact" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "This project advances AI-powered interactive entertainment by solving the sandbox problem "
                    "that has limited LLM-driven games. The Structured Chaos architecture is a reusable pattern "
                    "for any application requiring creative AI narration within rule-constrained systems: game "
                    "development, educational simulations, and training scenarios.\n\n"
                    "The platform can serve as an educational tool for language learning and narrative literacy. "
                    "The multiplayer co-op mode develops teamwork and critical thinking. The core architecture "
                    "(scenario schema, sandbox layers, procedural generation) will be open-sourced to allow "
                    "further research and development."
                )
                break
        break

# ── 12. RISK ANALYSIS ───────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Risk" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "Risk 1 — LLM generates invalid scenario JSON (Likelihood: Medium, Impact: High): "
                    "Mitigated by JSON schema validation and automatic re-prompt with specific errors. "
                    "Max 3 retries, then fallback to a pre-authored scenario.\n\n"
                    "Risk 2 — LLM sandbox breach via adversarial input (Likelihood: Low, Impact: High): "
                    "Mitigated by function call gating. The LLM has no API surface other than validated "
                    "functions. System prompt is a secondary measure only.\n\n"
                    "Risk 3 — Gemini API rate limits / costs (Likelihood: Medium, Impact: Medium): "
                    "Free tier (1500 req/day) is sufficient for development. Images cached per room. "
                    "Fallback placeholder if quota exceeded.\n\n"
                    "Risk 4 — Multiplayer latency (Likelihood: Low, Impact: Medium): Queue-based approach "
                    "tolerates async. No real-time twitch gameplay required. Acceptable wait is 3-5 seconds.\n\n"
                    "Risk 5 — Scope creep (Likelihood: Medium, Impact: Medium): WP1-WP3 alone form a "
                    "complete demonstrable project. WP4-WP5 are extensions."
                )
                break
        break

# ── 13. GANTT CHART ─────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh2" and "Gantt" in p.text:
        for j in range(i+1, min(i+6, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                # Clear text and insert chart image
                for run in np.runs:
                    run.text = ""
                if not np.runs:
                    np.add_run()
                np.runs[0].text = ""
                r = np.add_run()
                r.add_picture(_gantt_path, width=Inches(5.8))
                np.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Clear the MS Project line
                if j+1 < len(doc.paragraphs):
                    set_para(doc.paragraphs[j+1], "")
                break
        break

# ── 14. ECONOMICAL & ETHICAL ─────────────────────────────────────
# There are two Normal paragraphs under this heading
econ_texts = [
    "The entire backend and AI infrastructure runs on Google Cloud Platform (GCP). New GCP accounts receive "
    "$300 in free credits valid for 3 months, which is more than sufficient for the full semester of "
    "development and demonstration. LLM and image generation use Vertex AI (Gemini 2.5 Flash), covered under "
    "the same credit. The FastAPI backend is deployed on Google Cloud Run (serverless, scales to zero when idle). "
    "The frontend is hosted on Vercel free tier. No special hardware is required beyond standard development "
    "laptops. Total out-of-pocket cost for the semester is zero.",

    "Content Safety: The sandbox architecture limits what the LLM can narrate to pre-validated scenario "
    "elements, reducing the risk of harmful content. Content moderation may be added as a further safeguard.\n\n"
    "Intellectual Property: All content is AI-generated and clearly disclosed as such. No copyrighted fictional "
    "works, characters, or settings are used.\n\n"
    "Player Wellbeing: The platform will include session time indicators and will not employ dark patterns such "
    "as artificial urgency or manipulative reward schedules.\n\n"
    "Data Privacy: No personally identifiable information is stored beyond session duration. Users are informed "
    "that conversation logs are processed by Google Cloud under its data handling policies."
]

econ_idx = 0
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and "Economical" in p.text:
        count = 0
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                if count < len(econ_texts):
                    set_para(np, econ_texts[count])
                    count += 1
        break

# ── 15. REFERENCES ───────────────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "mh1" and "References" in p.text:
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            np = doc.paragraphs[j]
            if np.text.strip() and np.style.name == "Normal":
                set_para(np,
                    "[1] N. Montfort, Twisty Little Passages: An Approach to Interactive Fiction. MIT Press, 2003.\n"
                    "[2] A. Plotkin, The Inform Designer's Handbook, 2004. Available: http://inform-fiction.org\n"
                    "[3] P. Ammanabrolu et al., How to Avoid Being Eaten by a Grue, arXiv:2006.07409, 2020.\n"
                    "[4] N. Walton, AI Dungeon: Play any adventure you can imagine. Latitude, 2019. Available: https://aidungeon.io\n"
                    "[5] Google DeepMind, Gemini API: Function Calling. Google AI for Developers, 2024. Available: https://ai.google.dev/docs/function_calling\n"
                    "[6] N. Shaker, J. Togelius, M. Nelson, Procedural Content Generation in Games. Springer, 2016. Available: http://pcgbook.com"
                )
                break
        break

# ── Save ─────────────────────────────────────────────────────────
out = "/Users/batuhankaraman/demo491/COMP491_Proposal_FILLED.docx"
doc.save(out)
print(f"Saved: {out}")
